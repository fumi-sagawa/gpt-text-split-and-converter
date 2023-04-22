import openai
import tiktoken
import re
from docx import Document
from tiktoken.core import Encoding
from typing import List

GPT_MAX_TOKENS = 4096


def read_word_file(file_path: str) -> str:
    """
    Wordファイルを読み込み、テキストを返す関数。

    引数:
    - file_path (str): 読み込むWordファイルへのパス。

    戻り値:
    - str: Wordファイルから抽出されたテキスト。
    """
    doc = Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)


def split_text_at_max_tokens(text: str, max_tokens: int):
    """
    文章を最大トークン数ごとに切り詰める関数。

    引数：
    - text (str): 入力テキスト。
    - max_tokens (int): 分割する最大トークン数。

    戻り値：
    - List[str]: 最大トークン数で分割されたテキストチャンクのリスト。
    """

    subchunks = []

    encoding: Encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")
    start = 0
    while start < len(text):
        encoded_text = encoding.encode(text[start:])
        end = start + min(max_tokens, len(encoded_text))
        while end > start:
            tokens_to_encode = text[start:end]
            encoded_tokens = encoding.encode(tokens_to_encode)
            if len(encoded_tokens) <= max_tokens:
                break
            end -= 1

        subchunks.append(tokens_to_encode)
        start += end - start

    # Add the last subchunk to the list
    if start < len(text):
        subchunks.append(text[start:])

    return subchunks


def add_subchunks(subchunks: List[str], current_chunk: str, max_chunk_tokens: int, encoding: Encoding, system_prompt: str, order_prompt: str, chunks: List[str]) -> str:
    """
    現在のチャンクにサブチャンクを追加する関数。最大チャンクのトークン数を超えた場合は、現在のチャンクをリストに追加し新しいチャンクを開始する。

    引数：
    - subchunks (List[str]): 追加するサブチャンクのリスト。
    - current_chunk (str): 現在のチャンク。
    - max_chunk_tokens (int): 最大チャンクのトークン数。
    - encoding (Encoding): tiktokenのエンコーディングオブジェクト。
    - system_prompt (str): システムプロンプト。
    - order_prompt (str): オーダープロンプト。
    - chunks (List[str]): チャンクのリスト。

    戻り値：
    - str: 更新された（または新しい）現在のチャンク。
    """
    for subchunk in subchunks:
        new_chunk = current_chunk + "\n" + subchunk if current_chunk else subchunk
        new_chunk_tokens = encoding.encode(
            new_chunk) + encoding.encode(system_prompt) + encoding.encode(order_prompt)
        if len(new_chunk_tokens) <= max_chunk_tokens:

            current_chunk = new_chunk
        else:
            chunks.append(current_chunk)
            current_chunk = subchunk
    return current_chunk


def split_text(text: str, system_prompt: str, order_prompt: str, max_tokens: int = GPT_MAX_TOKENS) -> List[str]:
    """
    与えられたテキストを、システムプロンプトとオーダープロンプトを含めた最大トークン数を考慮し、チャンクに分割します。

    引数：
    - text (str): 分割するテキスト。
    - system_prompt (str): システムプロンプト。
    - order_prompt (str): オーダープロンプト。
    - max_tokens (int): GPTの最大トークン数。

    戻り値：
    - List[str]: 分割されたテキストチャンクのリスト。
    """
    encoding: Encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")
    system_prompt_tokens = encoding.encode(system_prompt)
    order_prompt_tokens = encoding.encode(order_prompt)
    prompt_tokens = len(system_prompt_tokens) + len(order_prompt_tokens)

    paragraphs = text.split("\n")
    chunks = []
    current_chunk = ""
    # トークン数の計算には、システムプロンプトとオーダープロンプトも含める必要があるため。
    # -100している理由は返却値に余裕を持たせるため
    max_chunk_tokens = (max_tokens - prompt_tokens) // 2 - 50
    print(f"max_chunk_tokens: {max_chunk_tokens}")

    for paragraph in paragraphs:
        paragraph_tokens = encoding.encode(paragraph)

        if len(paragraph_tokens) > max_chunk_tokens:
            subchunks = re.split(r'。|！|？|;', paragraph)
            if len(subchunks) == 1:
                subchunks = split_text_at_max_tokens(
                    paragraph, max_chunk_tokens)

            current_chunk = add_subchunks(
                subchunks, current_chunk, max_chunk_tokens, encoding, system_prompt, order_prompt, chunks)
        else:
            current_chunk = add_subchunks(
                [paragraph], current_chunk, max_chunk_tokens, encoding, system_prompt, order_prompt, chunks)

    if current_chunk:
        chunks.append(current_chunk)

    return chunks


def transform_text_with_gpt(chunk: str, system_prompt: str, order_prompt: str) -> str:
    """
    GPT APIを使用してテキストチャンクを変換（校正）する関数。

    引数：
    - chunk (str): 変換するテキストチャンク。
    - system_prompt (str): GPT APIに送信するシステムプロンプト。
    - order_prompt (str): GPT APIに送信するユーザープロンプト。
    - max_tokens (int): GPT APIの最大トークン数（デフォルトは4096）。

    戻り値：
    - str: GPT APIから受信した変換済み（校正済み）テキスト。
    """
    prompt = f'Here is your task.\n{order_prompt}\nPerform the task on the text given in "Input" and output only the string that follows in "Output".\n  input:{chunk}\n output: '

    messages = [
        {
            "role": "system",
            "content": system_prompt
        },
        {
            "role": "user",
            "content": prompt
        },
    ]

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages,
        max_tokens=GPT_MAX_TOKENS//2,
        temperature=0,
    )

    return response['choices'][0]['message']['content']


def correct_text(text_chunks: List[str], system_prompt: str, order_prompt: str) -> List[str]:
    """
     テキストチャンクをループしてGPT APIで校正し、校正結果を返す関数。

    引数：
    - text_chunks (List[str]): 校正するテキストチャンクのリスト。 - system_prompt (str): GPT APIに送信するシステムプロンプト。 - order_prompt (str): GPT APIに送信するユーザープロンプト。

    戻り値：
    - List[str]: 校正されたテキストチャンクのリスト。
    """
    corrected_chunks = []

    for chunk in text_chunks:
        corrected_chunk = transform_text_with_gpt(
            chunk, system_prompt, order_prompt)
        corrected_chunks.append(corrected_chunk)

    return corrected_chunks


def save_corrected_text_to_word_file(corrected_text: str, output_path: str):
    """ 
    校正されたテキストを新しいWordファイルに保存する関数。
    引数：
    - corrected_text (str): 校正されたテキスト。
    - output_path (str): 保存するWordファイルへのパス。

    戻り値：
    - なし
    """
    corrected_doc = Document()
    corrected_doc.add_paragraph(corrected_text)
    corrected_doc.save(output_path)


if __name__ == "__main__":
    openai.api_key = "APIキーをセット！"

    input_path = "./materials/from_file.docx"
    output_path = "./materials/to_file.docx"

    your_custom_system_prompt = 'You are a professional proofreader. With great attention to detail, excellent language and communication skills, flexibility, and a strong sense of responsibility, you edit and proofread various documents, providing high-quality texts.'
    your_order_prompt = "文章の誤字脱字を補完し句読点を補完してください。できるだけ原文のフォーマットを守ってください。"

    # 1. Read the Word file
    text = read_word_file(input_path)

    # 2. Split the text
    text_chunks = split_text(
        text, system_prompt=your_custom_system_prompt, order_prompt=your_order_prompt)

    # 3. Correct the text
    corrected_chunks = correct_text(
        text_chunks, your_custom_system_prompt, your_order_prompt)

    # 4. Join the corrected chunks
    corrected_text = "\n".join(corrected_chunks)

    # 5. Save the corrected text to a new Word file
    save_corrected_text_to_word_file(corrected_text, output_path)
